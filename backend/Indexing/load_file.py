#Indexing/load_file.py

import os
import logging
import re
from typing import List
from langchain_community.document_loaders import (
    TextLoader,
    PyPDFLoader,
    UnstructuredWordDocumentLoader,
    CSVLoader,
)
from config import NER_MODEL
from langchain.schema import Document
from transformers import pipeline  # HuggingFace NER 模型 / HuggingFace NER model
import pytesseract  # OCR 工具 / OCR tool
from pdf2image import convert_from_path
from PIL import Image
import fitz
from docx import Document as DocxDocument  # 使用 python-docx，解决 module 错误 / Use python-docx to resolve module issues

# 初始化日志
# Initialize logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

from config import NER_MODEL
from transformers import pipeline  # HuggingFace NER 模型 / HuggingFace NER model

# 初始化命名实体识别 (NER) 模型
# Initialize the Named Entity Recognition (NER) model
ner_pipeline = pipeline("ner", grouped_entities=True, model=NER_MODEL)


def apply_ner_to_text(text: str) -> str:
    """
    对文本内容应用命名实体识别 (NER) 并替换实体。
    Apply Named Entity Recognition (NER) to the text and replace entities.

    :param text: 输入文本 / Input text
    :return: 应用 NER 后的文本 / Text with NER applied
    """
    try:
        ner_results = ner_pipeline(text)
        for entity_group in ner_results:
            entity_text = entity_group["word"]
            entity_label = entity_group["entity_group"]
            text = text.replace(entity_text, f"[{entity_label}: {entity_text}]")
        return text
    except Exception as e:
        logging.error(f"NER failed. Error: {e}")
        return text


def clean_text(text: str) -> str:
    """
    对文本进行清理，移除多余嵌套内容或无效 NER 标签。
    Clean the text by removing redundant nested content or invalid NER tags.

    :param text: 输入文本 / Input text
    :return: 清理后的文本 / Cleaned text
    """
    text = re.sub(r"\[(\w+): \[(\w+): ([^\]]+)\]\]", r"[\1: \3]", text)
    text = re.sub(r"\[(\w+): \w\]\s?", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def ocr_pdf(file_path: str) -> List[Document]:
    """
    对 PDF 执行 OCR 提取文本。
    Perform OCR to extract text from a PDF.

    :param file_path: PDF 文件路径 / PDF file path
    :return: 包含提取内容的 `Document` 列表 / List of `Document` objects with extracted content
    """
    try:
        images = extract_images_from_pdf(file_path)
        text_list = [ocr_image(image) for image in images]
        return [Document(page_content=text, metadata={"source": file_path}) for text in text_list if text.strip()]
    except Exception as e:
        logging.error(f"OCR extraction failed: {file_path}. Error: {e}")
        return []


def extract_images_from_pdf(pdf_path: str) -> list:
    """
    从 PDF 文件中提取图片并保存。
    Extract images from a PDF file and save them.

    :param pdf_path: PDF 文件路径 / PDF file path
    :return: 提取后图片的路径列表 / List of paths to extracted images
    """
    try:
        images = []
        pdf_doc = fitz.open(pdf_path)
        for page_number in range(len(pdf_doc)):
            page = pdf_doc[page_number]
            image_list = page.get_images(full=True)  # 获取图像列表 / Get list of images

            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = pdf_doc.extract_image(xref)
                image_data = base_image["image"]
                ext = base_image["ext"]
                image_name = f"pdf_page_{page_number + 1}_{img_index + 1}.{ext}"
                image_path = os.path.join(os.getcwd(), image_name)
                with open(image_path, "wb") as img_file:
                    img_file.write(image_data)  # 保存图像文件 / Save the image file
                images.append(image_path)
        return images
    except Exception as e:
        logging.error(f"Error extracting images from PDF: {pdf_path}. Error: {e}")
        return []


def extract_images_from_docx(docx_path: str) -> List[str]:
    """
    从 DOCX 文件中提取图片。
    Extract images from a DOCX file.

    :param docx_path: DOCX 文件路径 / DOCX file path
    :return: 提取后图片的路径列表 / List of paths to extracted images
    """
    try:
        doc = DocxDocument(docx_path)
        images = []
        for rel in doc.part.rels.values():
            if "image" in rel.target_mode:
                image_name = os.path.basename(rel.target_ref)
                image_path = os.path.join(os.getcwd(), image_name)
                with open(image_path, "wb") as img_file:
                    img_file.write(rel.target_part.blob)
                images.append(image_path)
        return images
    except Exception as e:
        logging.error(f"Error extracting images from DOCX: {docx_path}. Error: {e}")
        return []


def ocr_image(image_path: str) -> str:
    """
    对图片执行 OCR 识别。
    Perform OCR recognition on an image.

    :param image_path: 图片路径 / Image file path
    :return: 提取的文本 / Extracted text
    """
    try:
        text = pytesseract.image_to_string(Image.open(image_path), lang="eng+chi_sim")
        return clean_text(text)
    except Exception as e:
        logging.error(f"Error during OCR for image: {image_path}. Error: {e}")
        return ""


def load_file(file_path: str) -> List[Document]:
    """
    加载文档（支持 TXT、PDF、DOCX、图片等格式）。
    Load documents (supports formats like TXT, PDF, DOCX, images, etc.).

    :param file_path: 文件路径 / File path
    :return: 加载的 `Document` 对象列表 / List of loaded `Document` objects
    """
    logging.info(f"Loading file: {file_path}")
    file_extension = os.path.splitext(file_path)[-1].lower()

    try:
        if file_extension == ".txt":
            loader = TextLoader(file_path, encoding="utf-8")
            documents = loader.load()

        elif file_extension == ".pdf":
            try:
                loader = PyPDFLoader(file_path)
                documents = loader.load()
            except Exception:
                logging.warning(f"PDF parse failed. Switching to OCR and image extraction: {file_path}")
                documents = []
                # OCR 文本提取 / OCR text extraction
                page_texts = ocr_pdf(file_path)
                documents.extend(page_texts)

                # 提取内嵌图片并 OCR / Extract embedded images and perform OCR
                image_paths = extract_images_from_pdf(file_path)
                for image_path in image_paths:
                    ocr_text = ocr_image(image_path)
                    if ocr_text:
                        documents.append(
                            Document(page_content=ocr_text, metadata={"source": file_path, "image": image_path})
                        )

        elif file_extension == ".docx":
            loader = UnstructuredWordDocumentLoader(file_path)
            documents = loader.load()
            image_paths = extract_images_from_docx(file_path)
            for image_path in image_paths:
                ocr_text = ocr_image(image_path)
                if ocr_text:
                    documents.append(
                        Document(page_content=ocr_text, metadata={"source": file_path, "image": image_path})
                    )

        elif file_extension in [".jpg", ".jpeg", ".png"]:
            ocr_text = ocr_image(file_path)
            documents = [Document(page_content=ocr_text, metadata={"source": file_path})]

        elif file_extension == ".csv":
            loader = CSVLoader(file_path)
            documents = loader.load()

        else:
            logging.warning(f"Unsupported file format: {file_extension}")
            return []

        for doc in documents:
            if isinstance(doc, Document):
                doc.metadata["source"] = file_path
                doc.page_content = apply_ner_to_text(doc.page_content)

        return documents

    except Exception as e:
        logging.error(f"Failed to load file: {file_path}. Error: {e}")
        return []


def traverse_directory(directory_path: str) -> List[Document]:
    """
    遍历目录并递归加载支持的文件格式。
    Traverse a directory and recursively load supported file formats.

    :param directory_path: 目录路径 / Directory path
    :return: 加载的 `Document` 对象列表 / List of loaded `Document` objects
    """
    all_documents = []
    try:
        for root, _, files in os.walk(directory_path):
            for file_name in files:
                file_path = os.path.join(root, file_name)
                documents = load_file(file_path)
                all_documents.extend(documents)
        all_documents = [doc for doc in all_documents if doc.page_content.strip()]
        return all_documents
    except Exception as e:
        logging.error(f"Failed to process directory: {directory_path}. Error: {e}")
        return []


# 测试代码
# Testing code
if __name__ == "__main__":
    logging.info("Starting document loading process...")
    directory_path = "../Data"  # 替换为实际的测试目录路径 / Replace with actual test directory path
    results = traverse_directory(directory_path)
    logging.info(f"Successfully loaded {len(results)} documents.")
    for doc in results[:5]:  # 仅输出部分结果 / Output partial results only
        print(doc.metadata)
        print(doc.page_content[:200])  # 显示前 200 个字符 / Display the first 200 characters
